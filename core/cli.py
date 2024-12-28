import argparse
from docx import Document
from docx.shared import Inches

from conf.config import Config


class GradBuilderCLI:
    """Класс для генерации диплома через CLI."""

    def __init__(self, config: Config):
        self.config = config
        self.language_manager = self.config.language_manager
        self.argument_parser = argparse.ArgumentParser(self.language_manager)

    def generate_diploma(self, args):
        """Генерирует диплом на основе аргументов."""
        # Обновляем язык, если указан
        self.language_manager.set_language(args.lang)

        # Создание документа
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(args.margin_left)
        section.right_margin = Inches(args.margin_right)
        section.top_margin = Inches(args.margin_top)
        section.bottom_margin = Inches(args.margin_bottom)

        # Титульный лист
        doc.add_heading(args.title, level=0)
        doc.add_paragraph(f"{self.language_manager.get('author')}: {args.author}")
        doc.add_paragraph(
            f"{self.language_manager.get('supervisor')}: {args.supervisor}"
        )
        doc.add_paragraph(f"{self.language_manager.get('year')}: {args.year}")

        # Сохранение файла
        doc.save(args.output)

        print(self.language_manager.get("status_success").format(path=args.output))

    def run(self):
        """Запускает CLI."""
        args = self.argument_parser.parse_arguments()
        self.generate_diploma(args)
